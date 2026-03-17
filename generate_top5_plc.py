import requests
import re
from datetime import date
from dotenv import load_dotenv
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

load_dotenv('C:/Users/sfaramarz/jill/.env')
jira_url = os.getenv('JIRA_BASE_URL')
username = os.getenv('JIRA_USERNAME')
token = os.getenv('JIRA_API_TOKEN')

# ── Step 1: Fetch all active LIGHTS parent tasks ──────────────────────────────
jql = ('project = LIGHTS AND summary ~ "Parent Task" '
       'AND status NOT IN ("Will Not Do", "Not Applicable") '
       'ORDER BY updated DESC')
resp = requests.get(f'{jira_url}/rest/api/2/search',
    params={'jql': jql, 'maxResults': 100,
            'fields': 'summary,status,assignee,comment,updated,subtasks'},
    auth=(username, token), headers={'Accept': 'application/json'})
resp.raise_for_status()
parent_tasks = resp.json()['issues']
print(f'Fetched {len(parent_tasks)} parent tasks')

# ── Step 1b: Override display names for ambiguous parent task summaries ───────
# Maps LIGHTS ticket key → correct "Tool Name Version" display string
TICKET_NAME_OVERRIDE = {
    'LIGHTS-499': 'NVRTX v5.7.3',          # "Post GDC Release"
    'LIGHTS-458': 'UE DLSS v8.5.0',         # "8.5.0"
    'LIGHTS-311': 'UE DLSS v8.4.0',         # "UE DLSS Plugin 8.4.0"
    'LIGHTS-260': 'RTX Remix v1.3.0',       # "v1.3.0"
    'LIGHTS-472': 'Comfy NV Video Prep',    # "1.4.0 ComfyUI AI Remaster Graph"
    'LIGHTS-465': 'IGI v1.5 GDC',           # "IGI 1.5 GDC Support"
    'LIGHTS-479': 'IGI v1.5',               # "NVIDIA In Game Inference SDK 1.5"
    'LIGHTS-177': 'RTX Remix v1.4',         # "RTX Remix 0.7.0" (active via REL-6)
    'LIGHTS-292': 'NVRTX v5.7',             # "NVIDIA RTX Unreal Engine NVRTX 5.7"
    'LIGHTS-353': 'RTXDI v3.0',
    'LIGHTS-272': 'RTXCR v1.1',
    'LIGHTS-519': 'RTXPT v1.8 / v3.0',
    'LIGHTS-506': 'Kokoro Plugin v1.0',
    'LIGHTS-432': 'RTXGI SDK v3.0',
    'LIGHTS-413': 'MegaGeometry SDK v2.0',
    'LIGHTS-505': 'NVRTX v5.7.3',
}

# ── Step 2: Known definitions ─────────────────────────────────────────────────
DEFINITIONS = {
    'rtxdi':              'Efficient dynamic lighting via ReSTIR spatiotemporal resampling',
    'nne':                'TensorRT AI/ML inference plugin for Unreal Engine',
    'ue nne':             'TensorRT AI/ML inference plugin for Unreal Engine',
    'dlss':               'DLSS SR, FG, MFG, RR, DLAA & Reflex for Unreal Engine',
    'ue dlss':            'DLSS SR, FG, MFG, RR, DLAA & Reflex for Unreal Engine',
    'igi':                'SDK for on-device AI inference in games',
    'in game inference':  'SDK for on-device AI inference in games',
    'megageometry':       'Ray tracing for massive Nanite scenes with fine-grain LOD',
    'rtxgi':              'Real-time ray traced global illumination SDK',
    'nvrtx':              'NVIDIA-maintained UE branch for RTX feature prototyping',
    'kokoro':             'Real-time AI text-to-speech for in-game use',
    'rtxpt':              'Open-source real-time path tracing SDK (DX12 / Vulkan)',
    'path tracing':       'Open-source real-time path tracing SDK (DX12 / Vulkan)',
    'comfy':              'ComfyUI AI remaster graph for RTX Remix asset prep',
    'comfyui':            'ComfyUI AI remaster graph for RTX Remix asset prep',
    'remix':              'AI-powered RTX remastering tool for classic games',
    'rtx remix':          'AI-powered RTX remastering tool for classic games',
    'frameview':          'Frame time & GPU/CPU performance measurement tool',
    'fvsdk':              'Frame time & GPU/CPU performance measurement tool',
    'rtxcr':              'Real-time character rendering SDK for skin, hair, and eyes',
    'character rendering': 'Real-time character rendering SDK for skin, hair, and eyes',
    'ace':                'AI-powered NPC animation and voice plugin for Unreal Engine',
    'avatar':             'Real-time digital human rendering and animation demo',
    'bonsai':             'NVIDIA Lightspeed Studios scene rendering demo',
    'kairos':             'Real-time ray-traced scene rendering SDK',
    'ai agent':           'Sample application for NVIDIA AI agent integration',
}

def get_definition(tool_name):
    name_lower = tool_name.lower()
    for key, defn in DEFINITIONS.items():
        if key in name_lower:
            return defn
    return 'NVIDIA developer SDK/plugin'

# ── Step 3: Parse tool name from LIGHTS summary ───────────────────────────────
def parse_tool_name(summary):
    # e.g. "L1 PLC Parent Task - [Kokoro Plugin ] [v1.0]"
    #      "L1 MVSB Parent Task - [RTXDI] [3.0]"
    #      "L1 Feature Level MVSB Parent Task - [Post GDC Release]"
    parts = re.findall(r'\[([^\]]+)\]', summary)
    if len(parts) >= 2:
        name = parts[0].strip()
        version = parts[1].strip()
        # Skip internal labels like "Post GDC Release", "Feature Name"
        skip = ['post gdc release', 'feature name', 'v1.3.0', '1.3.0', '8.5.0',
                'ue dlss plugin 8.5.0', 'ue dlss plugin 8.4.0', '8.4.0']
        if version.lower() not in skip:
            return f'{name} {version}'
        return name
    elif len(parts) == 1:
        return parts[0].strip()
    return summary

# ── Step 4: Map status to display value ──────────────────────────────────────
STATUS_MAP = {
    'backlog':      'To Start',
    'to do':        'To Start',
    'in progress':  'In Progress',
    'under review': 'In Progress',
    'signed-off':   'Done',
    'done':         'Done',
}
STATUS_SORT = {'Done': 0, 'In Progress': 1, 'To Start': 2}

def map_status(jira_status):
    return STATUS_MAP.get(jira_status.lower(), 'To Start')

# ── Step 5: Clean notes from Jira markup ─────────────────────────────────────
def clean_notes(text):
    if not text:
        return '—'
    # Remove Jira color markup
    text = re.sub(r'\{color[^}]*\}', '', text)
    text = re.sub(r'\{[^}]+\}', '', text)
    # Remove h1. h2. etc
    text = re.sub(r'h\d\.\s*', '', text)
    # Remove ~username~ mentions
    text = re.sub(r'\[~[^\]]+\]', '', text)
    # Remove bare URLs
    text = re.sub(r'https?://\S+', '', text)
    # Collapse whitespace
    text = re.sub(r'\s+', ' ', text).strip()
    # Truncate
    if len(text) > 180:
        text = text[:177] + '...'
    return text or '—'

# ── Step 5b: Release dates from REL project ──────────────────────────────────
rel_resp = requests.get(f'{jira_url}/rest/api/2/search',
    params={'jql': 'project = REL AND statusCategory != Done ORDER BY updated DESC',
            'maxResults': 50, 'fields': 'summary,comment,description,duedate'},
    auth=(username, token), headers={'Accept': 'application/json'})
rel_resp.raise_for_status()

# Extract release date from REL ticket comments/description
DATE_PATTERNS = [
    r'(?:release|target)[^\n]{0,30}?(\d{1,2}[\s/]\w+[\s/]\d{4})',
    r'(?:release|target)[^\n]{0,30}?((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{1,2},?\s+\d{4})',
    r'(?:release|target)[^\n]{0,30}?(\d{4}-\d{2}-\d{2})',
    r'((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\w*\s+\d{1,2},?\s+\d{4})',
]

def extract_date_from_text(text):
    if not text:
        return None
    for pat in DATE_PATTERNS:
        m = re.search(pat, text, re.IGNORECASE)
        if m:
            return m.group(1).strip()
    return None

# Build prog_key → release_date from REL tickets
REL_DATES = {}
PROG_FAMILY_REL = {
    'nvrtx': 'nvrtx', 'nvrtx': 'nvrtx',
    'dlss': 'ue dlss', 'ue dlss': 'ue dlss',
    'remix': 'remix',
    'frameview': 'frameview',
    'rtxgi': 'rtxgi',
    'megageometry': 'megageometry',
    'rtxpt': 'rtxpt',
    'kokoro': 'kokoro',
    'igi': 'igi',
}

for issue in rel_resp.json().get('issues', []):
    summary_lower = issue['fields']['summary'].lower()
    prog = next((v for k, v in PROG_FAMILY_REL.items() if k in summary_lower), None)
    if not prog or prog in REL_DATES:
        continue
    # Try duedate field first
    date_val = issue['fields'].get('duedate')
    if not date_val:
        # Try comments
        comments = issue['fields'].get('comment', {}).get('comments', [])
        for c in reversed(comments):
            date_val = extract_date_from_text(c.get('body', ''))
            if date_val:
                break
    if not date_val:
        date_val = extract_date_from_text(issue['fields'].get('description') or '')
    if date_val:
        REL_DATES[prog] = date_val

# Hardcode known dates not in REL comments
REL_DATES.setdefault('frameview',   'March 31, 2026')
REL_DATES.setdefault('remix',       'March 31, 2026 (may slip)')
REL_DATES.setdefault('nvrtx',       'GDC 2026')
REL_DATES.setdefault('igi',         'March 16, 2026')
REL_DATES.setdefault('ue dlss',     '8.5.1: Near-term | 8.6.0: TBD')
print('Release dates:', REL_DATES)

# ── Step 6: Build rows ────────────────────────────────────────────────────────
rows = []
seen_tools = set()

# Focus on recently active tasks (updated in last 6 months or still open)
from datetime import datetime, timezone
cutoff = datetime(2025, 1, 1, tzinfo=timezone.utc)  # only filter very old Done tasks

for task in parent_tasks:
    updated_str = task['fields'].get('updated', '')
    try:
        updated_dt = datetime.fromisoformat(updated_str.replace('Z', '+00:00'))
    except Exception:
        updated_dt = datetime.min.replace(tzinfo=timezone.utc)

    jira_status = task['fields']['status']['name']
    display_status = map_status(jira_status)

    # Never filter In Progress/To Start by date — only filter very old Done tasks
    SKIP_KEYWORDS = ['avatar demo', 'ace unreal', 'kairos', 'covert protocol',
                     'ai agent sample', 'bonsai', 'bridge initial']
    tool_lower = parse_tool_name(task['fields']['summary']).lower()
    if any(k in tool_lower for k in SKIP_KEYWORDS):
        continue
    if display_status == 'Done' and updated_dt < cutoff:
        continue

    tool_name = TICKET_NAME_OVERRIDE.get(task['key']) or parse_tool_name(task['fields']['summary'])

    # Map tool names to a program family key for dedup
    PROGRAM_FAMILY = {
        'nvrtx': 'nvrtx', 'nvidia rtx unreal': 'nvrtx',
        'ue dlss': 'ue dlss', 'dlss ue': 'ue dlss',
        'rtx remix': 'remix', 'remix': 'remix',
        'rtxdi': 'rtxdi',
        'rtxgi': 'rtxgi',
        'rtxpt': 'rtxpt', 'path tracing': 'rtxpt',
        'rtxcr': 'rtxcr', 'rtx character': 'rtxcr',
        'megageometry': 'megageometry', 'megageom': 'megageometry',
        'kokoro': 'kokoro',
        'frameview': 'frameview', 'fv': 'frameview',
        'igi': 'igi', 'in-game inference': 'igi', 'in game inference': 'igi',
        'ue nne': 'ue nne', 'nne plugin': 'ue nne',
        'comfy': 'comfy',
    }
    tool_lower = tool_name.lower()
    tool_key = next((v for k, v in PROGRAM_FAMILY.items() if tool_lower.startswith(k) or k in tool_lower), tool_lower[:12])
    if tool_key in seen_tools:
        # Replace existing Done entry with an In Progress one for the same program
        existing_idx = next((i for i, r in enumerate(rows) if r.get('prog_key') == tool_key), None)
        if existing_idx is not None and rows[existing_idx]['plc_status'] == 'Done' and display_status == 'In Progress':
            rows.pop(existing_idx)
            seen_tools.discard(tool_key)
        else:
            continue
    seen_tools.add(tool_key)

    # Notes: latest comment
    comments = task['fields'].get('comment', {}).get('comments', [])
    raw_notes = comments[-1]['body'] if comments else ''
    # Skip empty/whitespace-only or unicode-only comments
    stripped = re.sub(r'[\s\u2014\u2013\ufffd\u00a0]', '', raw_notes)
    notes = clean_notes(raw_notes) if stripped else '—'

    # Append open child ticket keys
    subtasks = task['fields'].get('subtasks', [])
    open_keys = []
    for s in subtasks:
        st_status = s['fields']['status']['statusCategory']['key']
        if st_status not in ('done', 'undefined'):
            open_keys.append(s['key'])
    if open_keys and notes == '—':
        notes = ', '.join(open_keys[:3])
    elif open_keys:
        notes = notes + '  ' + ', '.join(open_keys[:2])

    # Append parent task key as ticket reference if In Progress or To Start
    if display_status in ('In Progress', 'To Start') and task['key'] not in notes:
        if notes == '—':
            notes = task['key']
        else:
            notes = (notes.rstrip() + f'  {task["key"]}').strip()

    rows.append({
        'tool':         tool_name,
        'definition':   get_definition(tool_name),
        'plc_status':   display_status,
        'release_date': REL_DATES.get(tool_key, 'TBD'),
        'notes':        notes,
        'prog_key':     tool_key,
        'sort_key':     (STATUS_SORT.get(display_status, 3), -updated_dt.timestamp()),
    })

# ── Step 6b: Add programs not tracked via LIGHTS parent tasks ─────────────────
rows.append({
    'tool': 'FrameView v1.8.0',
    'definition': 'Frame time & GPU/CPU performance measurement tool',
    'plc_status': 'In Progress',
    'release_date': 'March 31, 2026',
    'notes': 'Coverity review, Release Review WIP, Legal done  FVSDK-20',
    'prog_key': 'frameview',
    'sort_key': (STATUS_SORT['In Progress'], -datetime(2026, 3, 13, tzinfo=timezone.utc).timestamp()),
})

rows.append({
    'tool': 'UE NNE Plugin',
    'definition': 'TensorRT AI/ML inference plugin for Unreal Engine',
    'plc_status': 'To Start',
    'release_date': 'TBD',
    'notes': '—',
    'prog_key': 'ue nne',
    'sort_key': (STATUS_SORT['To Start'], -datetime(2026, 3, 1, tzinfo=timezone.utc).timestamp()),
})

rows.append({
    'tool': 'UE DLSS v8.6.0',
    'definition': 'DLSS SR, FG, MFG, RR, DLAA & Reflex for Unreal Engine',
    'plc_status': 'To Start',
    'release_date': 'TBD',
    'notes': 'DLSS SDK dependency slipped; release date TBD  REL-2',
    'prog_key': 'ue dlss',
    'sort_key': (STATUS_SORT['To Start'], -datetime(2026, 3, 6, tzinfo=timezone.utc).timestamp()),
})

rows.sort(key=lambda r: r['sort_key'])
print(f'\nRows to include: {len(rows)}')
for r in rows:
    print(f'  [{r["plc_status"]:12}] {r["tool"]:35} | {r["notes"][:60]}')

# ── Step 7: Generate Word document ───────────────────────────────────────────
today_str     = date.today().strftime('%Y-%m-%d')
today_display = date.today().strftime('%d/%m/%Y')
desktop       = os.path.join(os.path.expanduser('~'), 'OneDrive - NVIDIA Corporation', 'Desktop')
output_path   = os.path.join(desktop, f'Top5_PLC_{today_str}.docx')

doc = Document()
for section in doc.sections:
    section.top_margin    = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)

def shade(cell, hex_fill):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_fill)
    tcPr.append(shd)

# Title
title = doc.add_heading('Top 5 Things', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
if title.runs:
    title.runs[0].font.color.rgb = RGBColor(0x76, 0xB9, 0x00)

# Mission
p = doc.add_paragraph()
r = p.add_run('Mission: ')
r.bold = True; r.font.size = Pt(10)
r2 = p.add_run(
    "Drive secure, compliant releases for NVIDIA's RTX developer tools and UE plugins "
    "(and more) through structured PLC governance."
)
r2.font.size = Pt(10)

# Dashboard line
p2 = doc.add_paragraph()
r3 = p2.add_run('Dashboard')
r3.bold = True; r3.font.size = Pt(10)
r3.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
r4 = p2.add_run(' contains live updates.')
r4.font.size = Pt(10)

doc.add_paragraph()

# Table — 5 columns
table = doc.add_table(rows=1, cols=5)
table.style = 'Table Grid'

col_widths = [Inches(1.4), Inches(2.4), Inches(1.0), Inches(1.2), Inches(2.0)]
for i, w in enumerate(col_widths):
    for cell in table.columns[i].cells:
        cell.width = w

hdr = table.rows[0].cells
for i, col_name in enumerate(['Tool', 'Definition', 'PLC Status', 'Release Date', 'Notes / Pending']):
    hdr[i].text = col_name
    run = hdr[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    shade(hdr[i], '404040')

STATUS_FILL = {'Done': 'D4EDDA', 'In Progress': 'D6E4FF', 'To Start': 'F4F5F7'}

for r in rows:
    row = table.add_row().cells
    row[0].text = r['tool']
    run = row[0].paragraphs[0].runs[0]
    run.font.size = Pt(9); run.bold = True

    row[1].text = r['definition']
    row[1].paragraphs[0].runs[0].font.size = Pt(9)

    row[2].text = r['plc_status']
    row[2].paragraphs[0].runs[0].font.size = Pt(9)
    shade(row[2], STATUS_FILL.get(r['plc_status'], 'FFFFFF'))

    row[3].text = r.get('release_date', 'TBD')
    row[3].paragraphs[0].runs[0].font.size = Pt(9)

    row[4].text = r['notes']
    row[4].paragraphs[0].runs[0].font.size = Pt(9)

doc.add_paragraph()

# Signature
doc.add_paragraph('Thanks,').runs[0].font.size = Pt(10)
doc.add_paragraph('Sherry Faramarz').runs[0].font.size = Pt(10)
doc.add_paragraph()

# Bcc
bcc_para = doc.add_paragraph()
bcc_run  = bcc_para.add_run(
    'Bcc: Jspitzer-staff, jpaul-org, GeForce-Devtech-Managers, DevStatus_UE, Producers, '
    'Keita Iida, Jaakko Haapasalo, KLM, Alex Dunn, John Spitzer, Jason Paul, '
    'Michael Songy, Nyle Usmani, Cem Cebenoyan'
)
bcc_run.italic = True
bcc_run.font.size = Pt(8)
bcc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.save(output_path)
print(f'\nSaved: {output_path}')
