"""
FrameView Status Report Generator
Produces a formatted .docx on the Desktop from assembled report data.

Expected inputs (set as variables before calling generate()):
  program_name      str   e.g. "FrameView"
  release_version   str   e.g. "1.8.0"
  release_date      str   e.g. "March 31, 2026"
  overall_status    str   "ON TRACK" | "AT RISK" | "OFF TRACK"
  executive_summary str   1-2 sentence summary
  bug_groups        dict  { "Group 1 — Performance & Capture Accuracy": [bug_dict, ...], ... }
  release_infra     dict  { "PBR #307197": "✅ Completed", ... }
  development_summary str
  planned_features  list  [str, ...]
"""

from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


STATUS_COLORS = {
    "ON TRACK":  RGBColor(0x00, 0x80, 0x00),
    "AT RISK":   RGBColor(0xFF, 0x99, 0x00),
    "OFF TRACK": RGBColor(0xCC, 0x00, 0x00),
}


def generate(program_name, release_version, release_date, overall_status,
             executive_summary, bug_groups, release_infra,
             development_summary, planned_features):

    today = date.today()
    today_str = today.strftime("%d_%m_%Y")
    today_display = today.strftime("%d/%m/%Y")
    today_formatted = today.strftime("%B %d, %Y")

    desktop = os.path.join(os.path.expanduser("~"), "OneDrive - NVIDIA Corporation", "Desktop")
    if not os.path.isdir(desktop):
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")
    output_path = os.path.join(desktop, f"{program_name}_Tool_SDK_Update_{today_str}.docx")

    # Handle locked file
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except PermissionError:
            output_path = output_path.replace(".docx", "_v2.docx")

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Title
    title = doc.add_heading(f"{program_name} Tool Update — Status Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(0x76, 0xB9, 0x00)

    # Subtitle
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Release {release_version}  |  Target: {release_date}  |  As of {today_formatted}")
    run.italic = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Overall Status
    status_para = doc.add_paragraph()
    status_para.add_run("Overall Status:  ").bold = True
    status_para.runs[0].font.size = Pt(12)
    status_run = status_para.add_run(overall_status)
    status_run.bold = True
    status_run.font.size = Pt(12)
    status_run.font.color.rgb = STATUS_COLORS.get(overall_status, RGBColor(0, 0, 0))

    # Executive summary — supports string or list of bullet strings
    if isinstance(executive_summary, list):
        for bullet in executive_summary:
            doc.add_paragraph(bullet, style="List Bullet").runs[0].font.size = Pt(10)
    else:
        doc.add_paragraph(executive_summary).runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # QA Bug Fix Status — flat table, no group headings
    doc.add_heading("QA Bug Fix Status", level=1)
    all_bugs = []
    if isinstance(bug_groups, dict):
        for bugs in bug_groups.values():
            all_bugs.extend(bugs)
    else:
        all_bugs = bug_groups
    if all_bugs:
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        for i, col in enumerate(["Bug ID", "Synopsis", "Status", "Engineer", "Last Updated", "Notes"]):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
            tc = hdr[i]._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "404040")
            tcPr.append(shd)
            hdr[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        for bug in all_bugs:
            row = table.add_row().cells
            for j, key in enumerate(["id", "synopsis", "status", "engineer", "last_updated", "notes"]):
                row[j].text = str(bug.get(key, ""))
                row[j].paragraphs[0].runs[0].font.size = Pt(9)
            if "P0" in bug.get("priority", ""):
                row[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        doc.add_paragraph()

    # Release Infrastructure
    doc.add_heading("Release Infrastructure", level=1)
    infra_table = doc.add_table(rows=1, cols=2)
    infra_table.style = "Table Grid"
    hdr = infra_table.rows[0].cells
    for i, col in enumerate(["Item", "Status"]):
        hdr[i].text = col
        hdr[i].paragraphs[0].runs[0].bold = True
        hdr[i].paragraphs[0].runs[0].font.size = Pt(9)
    for item, status in release_infra.items():
        row = infra_table.add_row().cells
        row[0].text = item
        row[1].text = status
        for cell in row:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
    doc.add_paragraph()

    # Development section
    doc.add_heading(f"{program_name} Development", level=1)
    doc.add_paragraph(development_summary).runs[0].font.size = Pt(10)
    if planned_features:
        doc.add_heading("Planned Features", level=2)
        for feat in planned_features:
            doc.add_paragraph(feat, style="List Bullet").runs[0].font.size = Pt(10)
    doc.add_paragraph()

    # Bcc & Signature
    bcc = doc.add_paragraph()
    bcc_run = bcc.add_run(
        "Bcc: Jspitzer-staff, jpaul-org, GeForce-Devtech-Managers, DevStatus_UE, Producers, "
        "Keita Iida, Jaakko Haapasalo, KLM, Alex Dunn, John Spitzer, Jason Paul, Michael Songy, "
        "Nyle Usmani, Cem Cebenoyan, frameview_devs"
    )
    bcc_run.italic = True
    bcc_run.font.size = Pt(8)
    bcc_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()
    doc.add_paragraph("Best Regards,\n\nSherry Faramarz").runs[0].font.size = Pt(10)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path
