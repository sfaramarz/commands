"""
PLC Top 5 Report Generator
Produces a single combined .docx on the Desktop for all LSS RTX Kit/Tools programs.

Expected input:
  rows  list of tuples:
    (tool_name: str, definition: str, release_date: str, plc_status: str, notes: str)

  Rows must be pre-sorted: Done -> In Progress -> To Start.
"""

from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# Status -> (text_color_hex, row_bg_hex)
STATUS_COLORS = {
    "Done":        ("008000", "E2F0D9"),
    "In Progress": ("BF8F00", "FFF2CC"),
    "To Start":    ("7F7F7F", "F2F2F2"),
}

HEADER_BG = "1F497D"


def _set_cell_shading(cell, color_hex):
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)


def _set_row_shading(row, color_hex):
    for cell in row.cells:
        _set_cell_shading(cell, color_hex)


def generate(rows, extra_note=None):
    """Generate the PLC Top 5 Word document.

    Args:
        rows: list of (tool_name, definition, release_date, plc_status, notes) tuples,
              pre-sorted Done -> In Progress -> To Start.
    Returns:
        output_path: str
    """
    today = date.today()
    today_str = today.strftime("%Y-%m-%d")

    desktop = os.path.join(
        os.path.expanduser("~"),
        "OneDrive - NVIDIA Corporation",
        "Desktop",
    )
    if not os.path.isdir(desktop):
        desktop = os.path.join(os.path.expanduser("~"), "Desktop")

    output_path = os.path.join(desktop, f"LSS_RTX_PLC_Top5_{today_str}.docx")

    # Remove stale file if possible
    if os.path.exists(output_path):
        try:
            os.remove(output_path)
        except PermissionError:
            output_path = output_path.replace(".docx", "_new.docx")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # --- Title ---
    title = doc.add_heading("Top 5 Things", level=1)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x76, 0xB9, 0x00)

    # --- Mission ---
    mission = doc.add_paragraph()
    m_bold = mission.add_run("Mission: ")
    m_bold.bold = True
    m_bold.font.size = Pt(10)
    m_text = mission.add_run(
        "Drive secure, compliant releases for NVIDIA\u2019s RTX developer "
        "tools and UE plugins (and more) through structured PLC governance."
    )
    m_text.italic = True
    m_text.font.size = Pt(10)

    # --- Dashboard link ---
    dash = doc.add_paragraph()
    d_link = dash.add_run("Dashboard")
    d_link.bold = True
    d_link.underline = True
    d_link.font.size = Pt(10)
    d_link.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    d_text = dash.add_run(" contains live updates.")
    d_text.italic = True
    d_text.font.size = Pt(10)

    # --- Extra note (optional) ---
    if extra_note:
        note_para = doc.add_paragraph()
        note_run = note_para.add_run(extra_note)
        note_run.italic = True
        note_run.font.size = Pt(9)
        note_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # --- Table ---
    table = doc.add_table(rows=len(rows) + 1, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_obj in table.rows:
        row_obj.cells[0].width = Cm(3.5)
        row_obj.cells[1].width = Cm(5.0)
        row_obj.cells[2].width = Cm(2.5)
        row_obj.cells[3].width = Cm(2.5)
        row_obj.cells[4].width = Cm(5.5)

    # Header
    header_row = table.rows[0]
    _set_row_shading(header_row, HEADER_BG)
    for i, h in enumerate(["Tool", "Definition", "Release Date", "PLC Status", "Notes / Pending"]):
        cell = header_row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, (tool, defn, rel_date, status, notes) in enumerate(rows):
        row_obj = table.rows[r_idx + 1]
        text_color, row_bg = STATUS_COLORS.get(status, ("000000", "FFFFFF"))
        _set_row_shading(row_obj, row_bg)

        # Tool (bold)
        c0 = row_obj.cells[0]
        c0.text = ""
        r0 = c0.paragraphs[0].add_run(tool)
        r0.font.size = Pt(10)
        r0.bold = True

        # Definition (italic)
        c1 = row_obj.cells[1]
        c1.text = ""
        r1 = c1.paragraphs[0].add_run(defn)
        r1.font.size = Pt(10)
        r1.italic = True

        # Release Date (centered)
        c2 = row_obj.cells[2]
        c2.text = ""
        r2 = c2.paragraphs[0].add_run(rel_date)
        r2.font.size = Pt(10)
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # PLC Status (centered, colored, bold)
        c3 = row_obj.cells[3]
        c3.text = ""
        r3 = c3.paragraphs[0].add_run(status)
        r3.font.size = Pt(10)
        r3.bold = True
        r3.font.color.rgb = RGBColor(
            int(text_color[0:2], 16),
            int(text_color[2:4], 16),
            int(text_color[4:6], 16),
        )
        c3.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Notes (9pt)
        c4 = row_obj.cells[4]
        c4.text = ""
        r4 = c4.paragraphs[0].add_run(notes)
        r4.font.size = Pt(9)

    doc.add_paragraph()

    # --- Sign-off ---
    thanks = doc.add_paragraph()
    thanks.add_run("Thanks,\nSherry Faramarz").font.size = Pt(10)

    doc.add_paragraph()

    # --- Bcc ---
    bcc = doc.add_paragraph()
    bcc_run = bcc.add_run(
        "Bcc: Jspitzer-staff, jpaul-org, GeForce-Devtech-Managers, "
        "DevStatus_UE, Producers, Keita Iida, Jaakko Haapasalo, KLM, "
        "Alex Dunn, John Spitzer, Jason Paul, Michael Songy, "
        "Nyle Usmani, Cem Cebenoyan"
    )
    bcc_run.italic = True
    bcc_run.font.size = Pt(8)
    bcc_run.font.color.rgb = RGBColor(0x7F, 0x7F, 0x7F)

    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path
